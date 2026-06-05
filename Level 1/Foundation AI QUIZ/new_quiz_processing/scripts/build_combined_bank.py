from __future__ import annotations

import difflib
import importlib.util
import json
import re
from pathlib import Path

try:
    from rapidfuzz import fuzz, process
except ImportError:  # pragma: no cover - fallback for machines without rapidfuzz
    fuzz = None
    process = None

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT.parent
OUT = ROOT / "outputs"
OUT.mkdir(parents=True, exist_ok=True)

SS6_SCRIPT = BASE / "quiz_extraction" / "scripts" / "build_study_guide.py"
SS6_FIGURES = BASE / "quiz_extraction" / "question_crops"
KORSOB_JSON = ROOT / "extracted" / "korsob_chud1_questions.json"
ANSWER_CAT_JSON = ROOT / "extracted" / "answer_by_category_questions.json"
OCR_DIR = ROOT / "extracted" / "100exam_ocr"
SS5_OCR_DIR = ROOT / "extracted" / "ss5_ocr"
EXAM02_OCR_DIR = ROOT / "extracted" / "exam02_ocr"
EXAM03_OCR_DIR = ROOT / "extracted" / "exam03_ocr"
TEST_KORKA_TITLES = ROOT / "extracted" / "test_korka_titles.json"

DOCX_OUT = OUT / "Foundation_AI_Combined_Question_Bank_compact.docx"
MD_OUT = OUT / "Foundation_AI_Combined_Question_Bank_compact.md"
DUP_JSON = OUT / "duplicate_report.json"
SOURCE_JSON = OUT / "source_inventory.json"
SS5_DUP_JSON = OUT / "ss5_potential_duplicate_report.json"
REF_DUP_JSON = OUT / "reference_potential_duplicate_report.json"
ALL_SOURCE_DUP_JSON = OUT / "all_source_duplicate_report.json"
ALL_SOURCE_DUP_THRESHOLD = 0.85

RED = RGBColor(192, 0, 0)
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 91, 91)
BLACK = RGBColor(0, 0, 0)


def load_ss6() -> list[dict]:
    spec = importlib.util.spec_from_file_location("ss6_guide", SS6_SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    items = []
    for q in mod.QUESTIONS:
        figures = []
        for fig in q.get("figures", []):
            fig_path = BASE / "quiz_extraction" / fig
            if not fig_path.exists():
                fig_path = SS6_FIGURES / f"q{q['n']:03d}" / fig
            if fig_path.exists():
                figures.append(str(fig_path.resolve()))
        items.append(
            {
                "kind": "verified",
                "source": "SuperAI SS6_Foundation AI QUIZ.pdf",
                "source_id": f"SS6-{q['n']:03d}",
                "source_number": q["n"],
                "question": q["question"],
                "options": [{"label": str(i + 1), "text": opt} for i, opt in enumerate(q["options"])],
                "answer": q["answer"],
                "answer_label": "",
                "explanation": q["explanation"],
                "note": q.get("note", ""),
                "figures": figures,
            }
        )
    return items


def source_error_note(q: dict) -> str:
    n = q["number"]
    if n == 35:
        return (
            "ตรวจแก้: source ไฮไลท์ข้อ a แต่ตาม REST ทั่วไป PUT/PATCH ใช้แก้ไขข้อมูลได้จริง "
            "ดังนั้นโจทย์ 'ข้อใดกล่าวผิด' ไม่มีตัวเลือกที่ผิดชัดเจนในชุดนี้"
        )
    if n == 47:
        return (
            "ตรวจแก้: source ไฮไลท์ train_size=0.20 ซึ่งให้ train 20% ไม่ใช่ 80%; "
            "คำสั่งที่ถูกควรใช้ model_selection, ลำดับคืนค่า X_train, X_test, y_train, y_test "
            "และกำหนด train_size=0.80 หรือ test_size=0.20"
        )
    return ""


def korsob_reason(q: dict) -> str:
    note = source_error_note(q)
    if note:
        return note
    question = q["question"]
    answer = q["answer"]
    lower = (question + " " + answer).lower()
    if "ถูกทุกข้อ" in answer:
        return "ทุกตัวเลือกเป็นข้อความที่ถูกต้องตามหลักของหัวข้อนี้ จึงต้องเลือก 'ถูกทุกข้อ' ไม่ใช่เลือกเพียงข้อย่อยข้อเดียว"
    if answer in {"TRUE", "FALSE"}:
        return f"ข้อความในโจทย์มีค่าความจริงเป็น {answer}; ให้จำแก่นิยามของหัวข้อนี้มากกว่าจำตำแหน่งตัวเลือก"
    patterns = [
        (["chmod", "permission"], "chmod 600 จำกัดสิทธิ์ไฟล์ให้เจ้าของอ่าน/เขียนได้ ช่วยให้ไฟล์ token/credential ปลอดภัยตามที่ Kaggle ต้องการ"),
        (["cbm"], "Condition-Based Maintenance เหมาะกับเครื่องจักรสำคัญ เพราะต้องลงทุนวัดสภาพและเฝ้าระวังเพื่อป้องกัน downtime"),
        (["word-embedding", "embedding"], "Skip-gram, CBOW, GloVe และ FastText เป็นวิธีสร้างเวกเตอร์แทนคำที่พบได้จริงใน NLP"),
        (["api", "201"], "HTTP 201 Created ใช้เมื่อ request สร้าง resource ใหม่สำเร็จ"),
        (["docker pull"], "docker pull คือคำสั่งดึง image จาก registry ลงมาใช้ในเครื่อง"),
        (["วิดีโอ", "เฟรม"], "การลบเฟรมก่อนหน้ากับเฟรมปัจจุบันเป็นวิธีพื้นฐานในการตรวจจับการเคลื่อนไหว"),
        (["โครโมโซม", "crossover"], "ลูกจาก crossover ต้องประกอบจาก segment ของพ่อแม่เดิม ไม่ใช่บิตที่เกิดขึ้นเองโดยไม่มีที่มา"),
        (["cv2.imread"], "cv2.imread() เป็นฟังก์ชันของ OpenCV สำหรับอ่านภาพจากไฟล์"),
        (["wireless", "plc"], "PLC สื่อสารผ่านสายไฟ จึงไม่ใช่เทคโนโลยี wireless"),
        (["tacotron", "fastpitch"], "FastPitch ใช้การทำนาย pitch/duration เพื่อควบคุม prosody และเร่งการสังเคราะห์เสียงต่างจาก Tacotron2"),
        (["lexitron"], "Lexitron เป็นพจนานุกรมอิเล็กทรอนิกส์ไทย-อังกฤษ/อังกฤษ-ไทย"),
        (["pos tagging"], "POS tagging คือการระบุชนิดคำ เช่น noun, verb, adjective ให้กับ token"),
        (["histogram equalization"], "Histogram equalization ปรับกระจายค่าความเข้มให้ contrast ของภาพดีขึ้น"),
        (["missing value", "k-nn"], "K-NN ใช้ข้อมูลเพื่อนบ้านที่คล้ายกันมาช่วยเติม missing value ได้ทั้งข้อมูลต่อเนื่องและข้อมูลกลุ่ม"),
        (["sampling", "spatial"], "การ sampling ภาพคือการ discretize ตำแหน่งเชิงพื้นที่ ส่วน quantization เป็นการ discretize ค่าความเข้ม"),
        (["database", "graph db", "document"], "ชนิดฐานข้อมูลต้องจับจากโมเดลข้อมูล: graph, document, key-value cache, RDBMS และ wide-column"),
        (["abugidas", "ภาษาไทย"], "ภาษาไทยเป็นระบบ abugida หรืออักษรเชิงพยางค์ที่พยัญชนะมีเสียงสระพื้นฐานและใช้เครื่องหมายสระกำกับ"),
        (["keypoint"], "ในการ matching keypoint ต้องกรอง match ที่ดี ไม่ใช้ทุก match โดยตรง เพราะ outlier ทำให้ transformation ผิดได้"),
        (["noise", "spatial filtering"], "Spatial filtering ปรับค่าพิกเซลจากบริเวณรอบข้าง จึงใช้ลด noise ของภาพได้"),
        (["stemming", "lemmatization"], "Stemming ตัดรูปคำแบบ heuristic ส่วน lemmatization ใช้ความรู้ทางภาษาเพื่อคืน lemma ที่ถูกต้องกว่า"),
        (["category comparison"], "Category comparison ใช้แสดงข้อมูลเพื่อเปรียบเทียบค่าระหว่างหมวดหมู่"),
        (["nvidia-smi"], "!nvidia-smi แสดงสถานะ GPU ที่ Colab จัดสรรให้ เช่น รุ่นและหน่วยความจำ"),
        (["train_test_split"], "train_test_split ต้อง import จาก sklearn.model_selection และกำหนดสัดส่วน train/test ให้ตรงกับโจทย์"),
        (["convolution"], "ขนาด output ของ convolution คำนวณจาก input, kernel, stride และ padding แล้วคูณด้วยจำนวน filter"),
        (["proof of stake", "pos"], "Proof of Stake ไม่ต้องใช้เครื่องขุดแข่ง hash แบบ Proof of Work จึงประหยัดพลังงานกว่า"),
        (["sql", "nosql", "json"], "JSON เป็นรูปแบบข้อมูล ส่วน SQL/NoSQL เป็นแนวทางฐานข้อมูลที่ใช้จัดเก็บ/สืบค้นข้อมูล"),
        (["bayesian"], "Bayesian update ปรับความเชื่อเดิมด้วยหลักฐานใหม่ผ่าน likelihood เพื่อได้ posterior"),
        (["gradient"], "การไล่ตาม gradient ปรับตัวแปรไปในทิศที่ทำให้ฟังก์ชันเพิ่มหรือลดตามเป้าหมาย optimization"),
        (["sha256"], "SHA-256 เป็น hash ทางเดียว ใช้ตรวจการเปลี่ยนแปลงข้อมูล ไม่ใช่การเข้ารหัสที่ถอดกลับได้"),
        (["anova"], "เมื่อ reject H0 ใน ANOVA แปลว่าอย่างน้อยหนึ่งกลุ่มมีค่าเฉลี่ยแตกต่างจากกลุ่มอื่น"),
        (["docker logs"], "docker logs ตามด้วย container id/name ใช้ดู log ของ container นั้น"),
        (["outlier"], "การเพิ่ม feature ที่อธิบายบริบท เช่น ขนาดเรือหรือสภาพอากาศ ช่วยแยก outlier ที่แท้จริงจากความต่างตามสถานการณ์"),
        (["value chain"], "value chain ของข้อมูลเน้นกิจกรรมสร้างคุณค่าจากข้อมูล ไม่ใช่งาน governance/engineering ที่เป็น support โดยตรงเสมอไป"),
        (["iot"], "ระบบ IoT ประกอบด้วยอุปกรณ์/ฮาร์ดแวร์ ซอฟต์แวร์ และการเชื่อมต่อเพื่อรับส่งข้อมูล"),
        (["decision tree", "entropy"], "Decision Tree ใช้เกณฑ์อย่าง entropy/information gain เพื่อเลือก feature สำหรับ split"),
        (["bleu"], "BLEU พิจารณา n-gram overlap ระหว่าง candidate กับ reference โดยค่า cumulative จะลดเมื่อ n-gram ยาวขึ้นและ match น้อยลง"),
        (["first-order", "∀"], "ประโยคตรรกะนี้ต้องใช้ universal quantifier กับ implication เพื่อบอกว่านักศึกษาทุกคนที่เรียน SuperAI แล้วเก่ง"),
        (["flatten"], "Flatten แปลง tensor/array หลายมิติให้เป็นเวกเตอร์หนึ่งมิติ ก่อนส่งเข้า layer แบบ fully connected"),
        (["covariance"], "covariance matrix เป็นเมทริกซ์จัตุรัส ขนาดเท่าจำนวนตัวแปรที่นำมาคำนวณ"),
        (["phoneme"], "TTS มักใช้ phoneme หรือหน่วยเสียงเป็น representation สำคัญก่อนสังเคราะห์เสียง"),
        (["perceptron", "xor"], "Perceptron ชั้นเดียวแก้ XOR ไม่ได้เพราะ XOR ไม่ linearly separable"),
        (["lie factor"], "Lie factor ใกล้ 1 แปลว่ากราฟไม่บิดเบือนข้อมูลมาก ช่วง 0.95-1.05 ถือว่ายอมรับได้"),
        (["degree of freedom"], "DOF คือจำนวนตัวแปรอิสระที่ยังเปลี่ยนค่าได้ภายใต้ constraint"),
        (["icp"], "Vanilla ICP ทำซ้ำสองแกนหลักคือจับคู่จุดที่สอดคล้องกันและคำนวณ alignment"),
        (["bootstrap"], "Bootstrap sampling คือการสุ่มตัวอย่างแบบใส่คืน ทำให้ตัวอย่างซ้ำกันได้"),
        (["asr"], "ASR โดยย่อแปลงสัญญาณเสียงผ่าน acoustic model แล้วใช้ language model ช่วยเลือกข้อความที่เป็นไปได้"),
        (["u-net", "segnet"], "U-Net มี skip connection ส่งข้อมูลจาก encoder ไป decoder ช่วยรักษารายละเอียดเชิงพื้นที่"),
        (["self-attention"], "Self-attention คำนวณความสัมพันธ์ระหว่างตำแหน่งต่าง ๆ ใน sequence เพื่อให้บริบทไหลถึงกัน"),
        (["lda topic"], "LDA topic modeling ใช้ค้นหา topic แฝง ไม่ได้ออกแบบมาเพื่อแก้ไวยากรณ์ของเอกสาร"),
        (["ica"], "ICA เน้นแยกแหล่งสัญญาณอิสระ ไม่ใช่วิธีลดมิติหลักแบบ PCA"),
        (["selection pressure"], "Selection pressure อธิบายแรงกดดันในการคัดเลือกที่ทำให้ solution ดีมีโอกาสสืบทอดมากขึ้น"),
        (["influencer"], "Exclusive Influencer ไม่ใช่ประเภทมาตรฐานทั่วไปเมื่อแบ่ง influencer ตามขนาด/บทบาท"),
        (["histogram"], "Histogram แสดงการกระจายความถี่ของค่า ไม่ใช่กราฟจำนวน object โดยตรง"),
        (["kaggle team"], "กฎ submission ของ Kaggle โดยทั่วไปผูกกับ competition/team ไม่ได้เพิ่มตามจำนวนสมาชิกแบบอิสระ"),
        (["intrinsic"], "Intrinsic parameter เป็นคุณลักษณะภายในกล้อง เช่น focal length และ principal point"),
        (["metamask"], "Seed phrase ใช้กู้คืนกระเป๋าใน wallet อื่นได้หากยังเก็บไว้ปลอดภัย"),
        (["opencv", "video"], "OpenCV อ่านวิดีโอได้จากหลายแหล่ง เช่น ไฟล์ กล้อง หรือ stream"),
        (["information extraction"], "Information extraction ดึง entity/relation/event จากข้อความ ส่วน image classification เป็นงาน vision คนละกลุ่ม"),
    ]
    for keys, reason in patterns:
        if all(k in lower for k in keys):
            return reason
    return f"ตัวเลือกนี้ตรงกับนิยามหรือคำสั่งที่ใช้จริงของหัวข้อในโจทย์ โดยตัวเลือกอื่นเป็นคนละหน้าที่หรือขัดกับเงื่อนไขที่ถาม"


def load_korsob() -> list[dict]:
    raw = json.loads(KORSOB_JSON.read_text(encoding="utf-8"))
    items = []
    for q in raw:
        correction = source_error_note(q)
        answer = q["answer"]
        answer_label = q["answer_label"]
        if q["number"] == 35:
            answer = "ไม่มีตัวเลือกที่ผิดชัดเจน; source ไฮไลท์ข้อ a แต่ PUT/PATCH ใช้แก้ไขข้อมูลได้จริง"
            answer_label = "ตรวจแก้"
        elif q["number"] == 47:
            answer = (
                "ไม่มีตัวเลือกถูกครบ; คำสั่งที่ถูกควรเป็น from sklearn.model_selection import train_test_split; "
                "X_train, X_test, y_train, y_test = train_test_split(X, y, train_size=0.80) หรือ test_size=0.20"
            )
            answer_label = "ตรวจแก้"
        items.append(
            {
                "kind": "verified",
                "source": "ข้อสอบชุด1.pdf",
                "source_id": f"K1-{q['number']:03d}",
                "source_number": q["number"],
                "question": q["question"],
                "options": q["options"],
                "answer": answer,
                "answer_label": answer_label,
                "explanation": korsob_reason(q),
                "note": correction,
                "figures": [],
                "source_red_answer": q["answer"],
                "source_red_label": q["answer_label"],
            }
        )
    return items


def answer_category_reason(q: dict) -> str:
    note = q.get("answer_note", "")
    if note:
        if "source skipped" in note:
            return "ไฟล์เฉลยแยกหมวดไม่ได้ให้เฉลยของข้อนี้ จึงเก็บไว้เป็นรายการต้องตรวจภาพ/แหล่งอื่นเพิ่ม"
        if "no highlight" in note:
            return "ไฟล์เฉลยไม่ได้ไฮไลท์ตัวเลือก แต่ Random Forest เป็น ensemble ที่ใช้แนวคิด bagging โดยสุ่มข้อมูล/feature เพื่อรวมผลจากหลาย tree"
        if "body text" in note:
            return "คำตอบของข้อนี้เป็นคำอธิบายที่ถูกไฮไลท์ไว้ในเนื้อหา ไม่ใช่ตัวเลือกแบบ numbered choice"
    answer = q.get("answer", "")
    if answer in {"True", "TRUE", "False", "FALSE"}:
        return f"ข้อความในโจทย์มีค่าความจริงเป็น {answer} ตามเฉลยที่ไฮไลท์ในไฟล์แยกหมวด"
    return "เฉลยอ้างอิงจากข้อความที่ถูกไฮไลท์ไว้ในไฟล์ [เฉลย] ข้อสอบ แยกหมวด.pdf"


def load_answer_category() -> list[dict]:
    if not ANSWER_CAT_JSON.exists():
        return []
    raw = json.loads(ANSWER_CAT_JSON.read_text(encoding="utf-8"))
    items = []
    for q in raw:
        figures = []
        for fig in q.get("figures", []):
            fig_path = ROOT / fig
            if fig_path.exists() and is_real_figure(fig_path):
                figures.append(str(fig_path.resolve()))
        items.append(
            {
                "kind": "verified",
                "source": "[เฉลย] ข้อสอบ แยกหมวด.pdf",
                "source_id": f"CAT-{q['number']:03d}",
                "source_number": q["number"],
                "question": q["question"],
                "options": q.get("options", []),
                "answer": q.get("answer", ""),
                "answer_label": q.get("answer_label", ""),
                "explanation": answer_category_reason(q),
                "note": f"{q.get('category', '')}. {q.get('answer_note', '')}".strip(". "),
                "figures": figures,
            }
        )
    return items


def is_real_figure(path: Path) -> bool:
    try:
        with Image.open(path) as img:
            width, height = img.size
    except Exception:
        return False
    return width >= 80 and height >= 40


def normalize(text: str) -> str:
    text = text.lower().replace("\u200b", "")
    text = re.sub(r"[\s\W_]+", "", text, flags=re.UNICODE)
    return text


def similarity(a: str, b: str) -> float:
    na = normalize(a)
    nb = normalize(b)
    if not na or not nb:
        return 0.0
    if fuzz is not None:
        return fuzz.ratio(na, nb) / 100.0
    return difflib.SequenceMatcher(None, na, nb).ratio()


def dedupe(items: list[dict]) -> tuple[list[dict], list[dict]]:
    unique: list[dict] = []
    duplicates: list[dict] = []
    for item in items:
        best = None
        best_score = 0.0
        for existing in unique:
            score = similarity(item["question"], existing["question"])
            if score > best_score:
                best_score = score
                best = existing
        if best is not None and best_score >= 0.93:
            dup = {
                "duplicate": item["source_id"],
                "duplicate_source": item["source"],
                "kept": best["source_id"],
                "kept_source": best["source"],
                "score": round(best_score, 3),
                "question": item["question"],
                "kept_question": best["question"],
            }
            duplicates.append(dup)
            best.setdefault("duplicate_sources", []).append(dup)
            for fig in item.get("figures", []):
                if fig not in best.setdefault("figures", []):
                    best["figures"].append(fig)
        else:
            item["unique_id"] = f"U{len(unique) + 1:03d}"
            item["duplicate_sources"] = []
            unique.append(item)
    return unique, duplicates


def load_ocr_items() -> list[dict]:
    items = []
    for n in range(1, 101):
        p = OCR_DIR / f"q{n:03d}.json"
        if not p.exists():
            continue
        q = json.loads(p.read_text(encoding="utf-8"))
        lines = q.get("lines", [])
        quality = "low" if len(lines) < 4 else "medium"
        items.append(
            {
                "source": "100Exam_Lv1.pdf",
                "source_id": f"100EX-{n:03d}",
                "number": n,
                "ocr_text": q.get("text", ""),
                "line_count": len(lines),
                "quality": quality,
                "crop": str((ROOT / q["crop"]).resolve()),
            }
        )
    return items


def ocr_question_hint(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    useful = []
    for line in lines:
        low = line.lower()
        if "ข้อที่" in line or "ข้าที่" in line or "ลบคำ" in line or "ลบค่า" in line:
            continue
        useful.append(line)
        if len(" ".join(useful)) > 160:
            break
    return " ".join(useful[:3]) or (lines[0] if lines else "")


def load_ss5_items(unique: list[dict]) -> tuple[list[dict], list[dict]]:
    items = []
    potential_dups = []
    for n in range(1, 101):
        p = SS5_OCR_DIR / f"q{n:03d}.json"
        if not p.exists():
            continue
        q = json.loads(p.read_text(encoding="utf-8"))
        text = q.get("text", "")
        hint = ocr_question_hint(text)
        best = None
        best_score = 0.0
        for existing in unique:
            score = max(similarity(hint, existing["question"]), similarity(text[:450], existing["question"]))
            if score > best_score:
                best_score = score
                best = existing
        status = "new_or_unmatched"
        if best is not None and best_score >= 0.90:
            status = "likely_duplicate"
            potential_dups.append(
                {
                    "ss5": f"SS5-{n:03d}",
                    "best_match": best["source_id"],
                    "best_unique_id": best["unique_id"],
                    "score": round(best_score, 3),
                    "ss5_hint": hint,
                    "matched_question": best["question"],
                }
            )
        items.append(
            {
                "source": "SuperAI SS5_Foundation AI QUIZ.pdf",
                "source_id": f"SS5-{n:03d}",
                "number": n,
                "ocr_text": text,
                "hint": hint,
                "line_count": len(q.get("lines", [])),
                "selected_index": q.get("selected_index"),
                "quality": "medium" if len(q.get("lines", [])) >= 5 else "low",
                "duplicate_status": status,
                "best_match": best["source_id"] if best is not None else "",
                "best_score": round(best_score, 3),
                "crop": str((ROOT / q["crop"]).resolve()),
            }
        )
    return items, potential_dups


def attach_duplicate_status(items: list[dict], unique: list[dict], threshold: float = 0.90) -> list[dict]:
    potential_dups = []
    for item in items:
        text = item.get("ocr_text", "")
        hint = item.get("hint") or ocr_question_hint(text)
        item["hint"] = hint
        best = None
        best_score = 0.0
        for existing in unique:
            score = max(similarity(hint, existing["question"]), similarity(text[:500], existing["question"]))
            if score > best_score:
                best_score = score
                best = existing
        item["duplicate_status"] = "new_or_unmatched"
        item["best_match"] = best["source_id"] if best is not None else ""
        item["best_score"] = round(best_score, 3)
        if best is not None and best_score >= threshold:
            item["duplicate_status"] = "likely_duplicate"
            potential_dups.append(
                {
                    "ref": item["source_id"],
                    "source": item["source"],
                    "best_match": best["source_id"],
                    "best_unique_id": best["unique_id"],
                    "score": round(best_score, 3),
                    "hint": hint,
                    "matched_question": best["question"],
                }
            )
    return potential_dups


def load_exam_items(exam: str, ocr_dir: Path, source_name: str, prefix: str, unique: list[dict]) -> tuple[list[dict], list[dict]]:
    items = []
    for n in range(1, 101):
        p = ocr_dir / f"q{n:03d}.json"
        if not p.exists():
            continue
        q = json.loads(p.read_text(encoding="utf-8"))
        selected = q.get("selected_index")
        if selected is not None and selected > 5:
            selected = None
        text = q.get("text", "")
        items.append(
            {
                "source": source_name,
                "source_id": f"{prefix}-{n:03d}",
                "number": n,
                "ocr_text": text,
                "hint": ocr_question_hint(text),
                "line_count": len(q.get("lines", [])),
                "selected_index": selected,
                "quality": "medium" if len(q.get("lines", [])) >= 5 else "low",
                "crop": str((ROOT / q["crop"]).resolve()),
            }
        )
    return items, attach_duplicate_status(items, unique)


def match_text(item: dict) -> str:
    if item.get("question"):
        return item.get("question", "")
    text = item.get("hint") or ocr_question_hint(item.get("ocr_text", ""))
    if len(text) < 80:
        text = f"{text} {item.get('ocr_text', '')[:200]}".strip()
    return text[:260]


def best_pool_match(query_norm: str, choices: list[str], pool: list[dict]) -> tuple[int | None, float]:
    if not query_norm or not choices:
        return None, 0.0
    if process is not None and fuzz is not None:
        match = process.extractOne(query_norm, choices, scorer=fuzz.ratio)
        if match is None:
            return None, 0.0
        return match[2], match[1] / 100.0

    best_idx = None
    best_score = 0.0
    for idx, choice in enumerate(choices):
        score = difflib.SequenceMatcher(None, query_norm, choice).ratio()
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx, best_score


def match_is_confident(query_norm: str, best_norm: str, score: float, threshold: float) -> bool:
    if score < threshold:
        return False
    if min(len(query_norm), len(best_norm)) < 28:
        return score >= 0.96
    return True


def compute_all_source_unique(
    unique: list[dict],
    source_groups: list[tuple[str, list[dict]]],
    threshold: float = ALL_SOURCE_DUP_THRESHOLD,
) -> tuple[int, list[dict], dict]:
    pool: list[dict] = []
    choices: list[str] = []
    for item in unique:
        text = match_text(item)
        norm_text = normalize(text)
        pool.append(
            {
                "source_id": item["source_id"],
                "source": item["source"],
                "text": text,
                "norm": norm_text,
            }
        )
        choices.append(norm_text)

    report: list[dict] = []
    summary: dict[str, dict[str, int]] = {}
    for source_name, items in source_groups:
        added = 0
        dup_count = 0
        for item in items:
            text = match_text(item)
            norm_text = normalize(text)
            best_idx, best_score = best_pool_match(norm_text, choices, pool)
            best = pool[best_idx] if best_idx is not None else None
            item["all_source_duplicate_status"] = "new_or_unmatched"
            item["all_source_best_match"] = best["source_id"] if best is not None else ""
            item["all_source_best_score"] = round(best_score, 3)
            if best is not None and match_is_confident(norm_text, best["norm"], best_score, threshold):
                dup_count += 1
                item["all_source_duplicate_status"] = "likely_duplicate"
                report.append(
                    {
                        "ref": item["source_id"],
                        "source": item["source"],
                        "best_match": best["source_id"],
                        "best_source": best["source"],
                        "score": round(best_score, 3),
                        "hint": text,
                        "matched_text": best["text"],
                    }
                )
            else:
                added += 1
                item["all_source_unique_id"] = f"AU{len(pool) + 1:03d}"
                pool.append(
                    {
                        "source_id": item["source_id"],
                        "source": item["source"],
                        "text": text,
                        "norm": norm_text,
                    }
                )
                choices.append(norm_text)
        summary[source_name] = {
            "new_or_unmatched": added,
            "likely_duplicate": dup_count,
        }
    return len(pool), report, summary


def set_font(run, size=None, color=None, bold=None, italic=None, name="Tahoma"):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def para_fmt(p, before=0, after=2, line=0.95, left=0, hanging=0, keep=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if left:
        fmt.left_indent = Inches(left)
    if hanging:
        fmt.first_line_indent = Inches(-hanging)
    fmt.keep_with_next = keep


def set_columns(section, num=2, space=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))


def configure_section(section, columns=1):
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    set_columns(section, columns, 300)


def add_run(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    para_fmt(p, before=6 if level == 1 else 4, after=3, line=1.0, keep=True)
    add_run(p, text, size=12 if level == 1 else 9.5, color=BLUE, bold=True)
    return p


def option_is_answer(item: dict, option: dict) -> bool:
    if item.get("answer_label") in {"ตรวจแก้", ""} and item.get("source") == "ข้อสอบชุด1.pdf":
        return option.get("is_answer", False) and not source_error_note({"number": item["source_number"], "question": item["question"], "answer": item.get("source_red_answer", "")})
    if option.get("is_answer"):
        return True
    ans = normalize(item.get("answer", ""))
    return bool(ans and normalize(option.get("text", "")) == ans)


def add_question_docx(doc, item: dict):
    p = doc.add_paragraph()
    para_fmt(p, before=2, after=1, line=0.92, keep=True)
    add_run(p, f"{item['unique_id']} [{item['source_id']}] ", size=7.7, color=BLUE, bold=True)
    add_run(p, item["question"], size=7.7, color=BLACK, bold=True)
    if item.get("duplicate_sources"):
        add_run(p, f"  (ซ้ำ/ใกล้ซ้ำ {len(item['duplicate_sources'])} source)", size=6.8, color=MUTED, italic=True)

    if item.get("figures"):
        for fig in item["figures"]:
            ip = doc.add_paragraph()
            para_fmt(ip, before=1, after=1, line=0.9)
            run = ip.add_run()
            run.add_picture(fig, width=Inches(3.05))

    for opt in item["options"]:
        op = doc.add_paragraph()
        para_fmt(op, before=0, after=0, line=0.9, left=0.12, hanging=0.12)
        is_ans = option_is_answer(item, opt)
        prefix = f"{opt.get('label', '')}. " if opt.get("label") else ""
        add_run(op, prefix, size=7.2, color=RED if is_ans else BLACK, bold=is_ans)
        add_run(op, opt.get("text", ""), size=7.2, color=RED if is_ans else BLACK, bold=is_ans)

    if item.get("answer_label") == "ตรวจแก้":
        ap = doc.add_paragraph()
        para_fmt(ap, before=0, after=0, line=0.9, left=0.12)
        add_run(ap, "เฉลยตรวจแก้: ", size=7.1, color=RED, bold=True)
        add_run(ap, item["answer"], size=7.1, color=RED, bold=True)
        if item.get("source_red_answer"):
            sp = doc.add_paragraph()
            para_fmt(sp, before=0, after=0, line=0.9, left=0.12)
            add_run(sp, f"source สีแดงเดิม: {item.get('source_red_label')}. {item.get('source_red_answer')}", size=6.6, color=MUTED, italic=True)
    elif item.get("answer_label") in {"text", "unknown"} or not item.get("options"):
        ap = doc.add_paragraph()
        para_fmt(ap, before=0, after=0, line=0.9, left=0.12)
        add_run(ap, "เฉลย: ", size=7.1, color=RED, bold=True)
        add_run(ap, item.get("answer", ""), size=7.1, color=RED, bold=True)

    ep = doc.add_paragraph()
    para_fmt(ep, before=0, after=2, line=0.9, left=0.12)
    add_run(ep, "เหตุผล: ", size=6.9, color=MUTED, bold=True)
    add_run(ep, item["explanation"], size=6.9, color=MUTED)
    if item.get("note") and item.get("answer_label") != "ตรวจแก้":
        np = doc.add_paragraph()
        para_fmt(np, before=0, after=2, line=0.9, left=0.12)
        add_run(np, "หมายเหตุ: ", size=6.8, color=RED, bold=True)
        add_run(np, item["note"], size=6.8, color=MUTED)


def build_docx(
    unique: list[dict],
    duplicates: list[dict],
    ss5_items: list[dict],
    ss5_dups: list[dict],
    exam02_items: list[dict],
    exam02_dups: list[dict],
    exam03_items: list[dict],
    exam03_dups: list[dict],
    ocr_items: list[dict],
    ocr_dups: list[dict],
    test_titles: list[dict],
    all_source_unique: int,
    all_source_dups: list[dict],
    all_source_summary: dict,
):
    doc = Document()
    configure_section(doc.sections[0], columns=1)
    styles = doc.styles
    styles["Normal"].font.name = "Tahoma"
    styles["Normal"].font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(title, after=4, line=1.0)
    add_run(title, "Foundation AI Combined Question Bank", size=15, color=BLUE, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(sub, after=6, line=1.0)
    add_run(sub, "Compact DOCX: เฉลยสีแดง + เหตุผล + duplicate/source audit", size=8, color=MUTED)

    summary = doc.add_paragraph()
    para_fmt(summary, after=3, line=1.0)
    add_run(summary, f"Main verified/deduped bank: {len(unique)} ข้อ จาก SS6 + ข้อสอบชุด1 + [เฉลย] แยกหมวด; duplicate/near-duplicate ที่แยกออก {len(duplicates)} รายการ. ", size=7.8)
    add_run(summary, f"All-source likely-unique estimate หลังรวม reference OCR appendices และ cross-source dedupe: {all_source_unique}+ ข้อ. Reference image-only: SS5 {len(ss5_items)}, Exam02 {len(exam02_items)}, Exam03 {len(exam03_items)}, 100Exam {len(ocr_items)}.", size=7.8, color=RED)

    add_heading(doc, "Source Audit", 1)
    audit_lines = [
        "SuperAI SS6_Foundation AI QUIZ.pdf: image long screenshot 100 ข้อ, ใช้ไฟล์ถอด/เฉลยที่ตรวจแก้ไว้ก่อนหน้าและภาพประกอบ q020/q074/q094/q095",
        "SuperAI SS5_Foundation AI QUIZ.pdf: image long screenshot 100 ข้อ, crop ครบ 100, OCR index ครบ 100, selected answer จากวงกลมในภาพ 99/100; ใช้เป็น appendix พร้อม duplicate status",
        "ข้อสอบชุด1.pdf: text PDF 91 ข้อ, ดึงเฉลยจากตัวอักษรสีแดง แล้วตรวจแก้ source-error ที่พบชัดเจนในข้อ K1-035 และ K1-047",
        "[เฉลย] ข้อสอบ แยกหมวด.pdf: text PDF 100 ข้อ, ดึงเฉลยจาก highlight สีเหลือง/ส้มและเพิ่มเข้า main verified bank พร้อมภาพประกอบ 4 รูป",
        "Exam02.pdf: image-block PDF 100 ข้อ, crop ครบ 100, OCR index ครบ 100, selected answer จากวงกลมในภาพ; ใช้เป็น appendix พร้อม duplicate status",
        "Exam03.pdf: image-block PDF 100 ข้อ, crop ครบ 100, OCR index ครบ 100, selected answer บางข้อจับได้จากภาพ; ใช้เป็น appendix พร้อม duplicate status",
        "100Exam_Lv1.pdf: image-only 100 ข้อ, crop high-res ครบ 100 และ OCR draft ครบ 100 แต่ยังไม่ใช้เป็นเฉลยแดง",
        "test_korka_compressed.pdf: screenshot/text-note 105 title lines ใช้เป็น source index สำหรับเทียบหัวข้อและรอตรวจมือ",
        "Quiz-92 score.docx/pdf: answer-audit resource โดยเฉพาะข้อ ROUGE/F-measure ใช้ประกอบหมายเหตุในชุด SS6",
    ]
    for line in audit_lines:
        p = doc.add_paragraph()
        para_fmt(p, after=1, line=0.95)
        add_run(p, line, size=7.4)

    doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(doc.sections[-1], columns=2)
    add_heading(doc, "Verified Deduped Bank", 1)
    for item in unique:
        add_question_docx(doc, item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(doc.sections[-1], columns=1)
    add_heading(doc, "Duplicate / Near-Duplicate Map", 1)
    if duplicates:
        for dup in duplicates:
            p = doc.add_paragraph()
            para_fmt(p, after=1, line=0.95)
            add_run(p, f"{dup['duplicate']} -> {dup['kept']} ", size=7.2, color=BLUE, bold=True)
            add_run(p, f"(score {dup['score']}) ", size=7.1, color=MUTED)
            add_run(p, re.sub(r"\s+", " ", dup["question"]).strip(), size=7.1)
    else:
        p = doc.add_paragraph()
        add_run(p, "ไม่พบข้อซ้ำแบบข้อความใกล้เคียงตาม threshold 0.93", size=7.4)

    add_heading(doc, "Reference Cross-Source Duplicate Estimate", 1)
    p = doc.add_paragraph()
    para_fmt(p, after=2, line=0.95)
    add_run(p, f"ใช้ OCR similarity threshold {ALL_SOURCE_DUP_THRESHOLD:.2f} เพื่อแยก likely duplicate ระหว่าง SS5/Exam02/Exam03/100Exam กับ main bank และกับ reference ที่เพิ่มก่อนหน้า; test_korka เป็น title index จึงไม่นับเป็น full question.", size=7.2, color=RED, bold=True)
    p = doc.add_paragraph()
    para_fmt(p, after=2, line=0.95)
    summary_bits = [f"{name}: +{stats['new_or_unmatched']} / dup {stats['likely_duplicate']}" for name, stats in all_source_summary.items()]
    add_run(p, "; ".join(summary_bits), size=7.1, color=MUTED)
    for dup in all_source_dups:
        p = doc.add_paragraph()
        para_fmt(p, after=1, line=0.95)
        add_run(p, f"{dup['ref']} -> {dup['best_match']} ", size=7.0, color=BLUE, bold=True)
        add_run(p, f"(score {dup['score']}) ", size=6.8, color=MUTED)
        hint = re.sub(r"\s+", " ", dup["hint"]).strip()
        add_run(p, hint[:220] + ("..." if len(hint) > 220 else ""), size=6.8)

    add_heading(doc, "SS5 Selected-Answer OCR/Crop Appendix", 1)
    p = doc.add_paragraph()
    para_fmt(p, after=3, line=0.95)
    add_run(p, "ส่วนนี้เพิ่มจาก SuperAI SS5_Foundation AI QUIZ.pdf ซึ่งเป็น image-only: OCR ใช้ค้นหา/เทียบซ้ำ, selected answer มาจากวงกลมทึบในภาพ, crop อยู่ใน ss5/question_crops. ยังไม่ถือเป็นเฉลยตรวจทฤษฎีละเอียดแบบ main bank.", size=7.4, color=RED, bold=True)
    for item in ss5_items:
        p = doc.add_paragraph()
        para_fmt(p, after=1, line=0.9)
        selected = item["selected_index"] if item["selected_index"] is not None else "unknown"
        status = "อาจซ้ำ" if item["duplicate_status"] == "likely_duplicate" else "เพิ่มเป็นรายการใหม่/ยังไม่ match"
        add_run(p, f"{item['source_id']} sel={selected} {status} ", size=6.9, color=BLUE, bold=True)
        if item["duplicate_status"] == "likely_duplicate":
            add_run(p, f"(match {item['best_match']} score {item['best_score']}) ", size=6.6, color=MUTED)
        if item.get("all_source_duplicate_status") == "likely_duplicate":
            add_run(p, f"(all-source match {item.get('all_source_best_match')} score {item.get('all_source_best_score')}) ", size=6.6, color=RED)
        text = re.sub(r"\s+", " ", item["ocr_text"]).strip()
        add_run(p, text[:430] + ("..." if len(text) > 430 else ""), size=6.7)

    def add_reference_appendix(title: str, intro: str, items: list[dict]) -> None:
        add_heading(doc, title, 1)
        p = doc.add_paragraph()
        para_fmt(p, after=3, line=0.95)
        add_run(p, intro, size=7.4, color=RED, bold=True)
        for item in items:
            p = doc.add_paragraph()
            para_fmt(p, after=1, line=0.9)
            selected = item.get("selected_index")
            selected_text = selected if selected is not None else "unknown"
            status = "อาจซ้ำ" if item.get("duplicate_status") == "likely_duplicate" else "เพิ่มเป็นรายการใหม่/ยังไม่ match"
            add_run(p, f"{item['source_id']} sel={selected_text} {status} ", size=6.9, color=BLUE, bold=True)
            if item.get("duplicate_status") == "likely_duplicate":
                add_run(p, f"(match {item.get('best_match')} score {item.get('best_score')}) ", size=6.6, color=MUTED)
            if item.get("all_source_duplicate_status") == "likely_duplicate":
                add_run(p, f"(all-source match {item.get('all_source_best_match')} score {item.get('all_source_best_score')}) ", size=6.6, color=RED)
            text = re.sub(r"\s+", " ", item.get("ocr_text", "")).strip()
            add_run(p, text[:430] + ("..." if len(text) > 430 else ""), size=6.7)

    add_reference_appendix(
        "Exam02 Selected-Answer OCR/Crop Appendix",
        "Exam02 เป็น image-block PDF 100 ข้อ: OCR ใช้ค้นหา/เทียบซ้ำ, selected answer มาจากวงกลมทึบในภาพ; ยังไม่ถือเป็นเฉลยตรวจทฤษฎีละเอียดแบบ main bank.",
        exam02_items,
    )
    add_reference_appendix(
        "Exam03 Selected-Answer OCR/Crop Appendix",
        "Exam03 เป็น image-block PDF 100 ข้อ: OCR ใช้ค้นหา/เทียบซ้ำ, selected answer บางข้อจับได้จากวงกลมในภาพ; unknown หมายถึง detector ไม่มั่นใจ จึงไม่เดา.",
        exam03_items,
    )

    add_heading(doc, "100Exam OCR/Crop Appendix", 1)
    p = doc.add_paragraph()
    para_fmt(p, after=3, line=0.95)
    add_run(p, "ส่วนนี้เป็น OCR draft จากไฟล์ image-only เพื่อใช้ค้นหาและตรวจมือ คู่กับ crop ในโฟลเดอร์ 100exam_hires_crops; ยังไม่ถือเป็น answer key.", size=7.4, color=RED, bold=True)
    for item in ocr_items:
        p = doc.add_paragraph()
        para_fmt(p, after=1, line=0.9)
        status = "อาจซ้ำ" if item.get("duplicate_status") == "likely_duplicate" else "เพิ่มเป็นรายการใหม่/ยังไม่ match"
        add_run(p, f"{item['source_id']} ({item['quality']}, {item['line_count']} lines, {status}): ", size=6.9, color=BLUE, bold=True)
        if item.get("duplicate_status") == "likely_duplicate":
            add_run(p, f"(match {item.get('best_match')} score {item.get('best_score')}) ", size=6.6, color=MUTED)
        if item.get("all_source_duplicate_status") == "likely_duplicate":
            add_run(p, f"(all-source match {item.get('all_source_best_match')} score {item.get('all_source_best_score')}) ", size=6.6, color=RED)
        text = re.sub(r"\s+", " ", item["ocr_text"]).strip()
        add_run(p, text[:450] + ("..." if len(text) > 450 else ""), size=6.7)

    add_heading(doc, "test_korka Title Index", 1)
    for i, title in enumerate(test_titles, start=1):
        p = doc.add_paragraph()
        para_fmt(p, after=0, line=0.9)
        add_run(p, f"TK-{i:03d} p{title['page']}: ", size=6.8, color=BLUE, bold=True)
        add_run(p, title["title"], size=6.8)

    doc.save(DOCX_OUT)


def md_escape(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def build_markdown(
    unique: list[dict],
    duplicates: list[dict],
    ss5_items: list[dict],
    ss5_dups: list[dict],
    exam02_items: list[dict],
    exam02_dups: list[dict],
    exam03_items: list[dict],
    exam03_dups: list[dict],
    ocr_items: list[dict],
    ocr_dups: list[dict],
    test_titles: list[dict],
    all_source_unique: int,
    all_source_dups: list[dict],
    all_source_summary: dict,
):
    lines = []
    lines.append("# Foundation AI Combined Question Bank")
    lines.append("")
    lines.append(f"- Main verified/deduped bank: {len(unique)} ข้อ")
    lines.append(f"- Duplicate/near-duplicate separated: {len(duplicates)} รายการ")
    lines.append(f"- SS5 image-only appendix: {len(ss5_items)} ข้อ; likely duplicates by OCR: {len(ss5_dups)} รายการ")
    lines.append(f"- Exam02 image-block appendix: {len(exam02_items)} ข้อ; likely duplicates by OCR: {len(exam02_dups)} รายการ")
    lines.append(f"- Exam03 image-block appendix: {len(exam03_items)} ข้อ; likely duplicates by OCR: {len(exam03_dups)} รายการ")
    lines.append(f"- 100Exam image-only appendix: {len(ocr_items)} ข้อ; likely duplicates by OCR: {len(ocr_dups)} รายการ")
    lines.append(f"- All-source likely-unique estimate after OCR cross-source dedupe: {all_source_unique}+ ข้อ")
    lines.append(f"- All-source OCR duplicate threshold: {ALL_SOURCE_DUP_THRESHOLD:.2f}")
    lines.append("- Correct answers are marked with `<span style=\"color:red\">...` in this Markdown and red text in DOCX.")
    lines.append("- 100Exam_Lv1 image-only items are included as OCR/crop appendix, not merged as verified answer key.")
    lines.append("")
    lines.append("## Verified Deduped Bank")
    for item in unique:
        lines.append("")
        lines.append(f"### {item['unique_id']} [{item['source_id']}]")
        lines.append(f"**คำถาม:** {md_escape(item['question'])}")
        if item.get("duplicate_sources"):
            lines.append(f"**ซ้ำ/ใกล้ซ้ำ:** {len(item['duplicate_sources'])} source")
        for fig in item.get("figures", []):
            rel = Path(fig).relative_to(BASE)
            lines.append(f"![figure]({rel.as_posix()})")
        for opt in item["options"]:
            text = f"{opt.get('label', '')}. {md_escape(opt.get('text', ''))}".strip()
            if option_is_answer(item, opt):
                lines.append(f"- <span style=\"color:red\">**{text}**</span>")
            else:
                lines.append(f"- {text}")
        if item.get("answer_label") == "ตรวจแก้":
            lines.append(f"- <span style=\"color:red\">**เฉลยตรวจแก้:** {md_escape(item['answer'])}</span>")
            if item.get("source_red_answer"):
                lines.append(f"- source สีแดงเดิม: {item.get('source_red_label')}. {md_escape(item.get('source_red_answer', ''))}")
        elif item.get("answer_label") in {"text", "unknown"} or not item.get("options"):
            lines.append(f"- <span style=\"color:red\">**เฉลย:** {md_escape(item.get('answer', ''))}</span>")
        lines.append(f"**เหตุผล:** {md_escape(item['explanation'])}")
        if item.get("note") and item.get("answer_label") != "ตรวจแก้":
            lines.append(f"**หมายเหตุ:** {md_escape(item['note'])}")

    lines.append("")
    lines.append("## Duplicate / Near-Duplicate Map")
    for dup in duplicates:
        lines.append(f"- {dup['duplicate']} -> {dup['kept']} (score {dup['score']}): {md_escape(dup['question'])}")

    lines.append("")
    lines.append("## Reference Cross-Source Duplicate Estimate")
    for name, stats in all_source_summary.items():
        lines.append(f"- {name}: +{stats['new_or_unmatched']} / likely duplicate {stats['likely_duplicate']}")
    for dup in all_source_dups:
        lines.append(f"- {dup['ref']} -> {dup['best_match']} (score {dup['score']}): {md_escape(dup['hint'])}")

    lines.append("")
    lines.append("## SS5 Selected-Answer OCR/Crop Appendix")
    lines.append("OCR draft + selected answer detected from filled radio circle. Use crops in `new_quiz_processing/ss5/question_crops` to verify.")
    for item in ss5_items:
        rel = Path(item["crop"]).relative_to(BASE)
        selected = item["selected_index"] if item["selected_index"] is not None else "unknown"
        lines.append("")
        lines.append(f"### {item['source_id']} selected={selected} duplicate_status={item['duplicate_status']} score={item['best_score']}")
        if item["duplicate_status"] == "likely_duplicate":
            lines.append(f"Likely duplicate: `{item['best_match']}`")
        if item.get("all_source_duplicate_status") == "likely_duplicate":
            lines.append(f"All-source likely duplicate: `{item.get('all_source_best_match')}` score={item.get('all_source_best_score')}")
        lines.append(f"Crop: `{rel.as_posix()}`")
        lines.append("```")
        lines.append(item["ocr_text"])
        lines.append("```")

    def append_ref_section(title: str, items: list[dict]) -> None:
        lines.append("")
        lines.append(f"## {title}")
        for item in items:
            rel = Path(item["crop"]).relative_to(BASE)
            selected = item.get("selected_index") if item.get("selected_index") is not None else "unknown"
            lines.append("")
            lines.append(f"### {item['source_id']} selected={selected} duplicate_status={item.get('duplicate_status')} score={item.get('best_score')}")
            if item.get("duplicate_status") == "likely_duplicate":
                lines.append(f"Likely duplicate: `{item.get('best_match')}`")
            if item.get("all_source_duplicate_status") == "likely_duplicate":
                lines.append(f"All-source likely duplicate: `{item.get('all_source_best_match')}` score={item.get('all_source_best_score')}")
            lines.append(f"Crop: `{rel.as_posix()}`")
            lines.append("```")
            lines.append(item.get("ocr_text", ""))
            lines.append("```")

    append_ref_section("Exam02 Selected-Answer OCR/Crop Appendix", exam02_items)
    append_ref_section("Exam03 Selected-Answer OCR/Crop Appendix", exam03_items)

    lines.append("")
    lines.append("## 100Exam OCR/Crop Appendix")
    lines.append("OCR draft only. Use high-res crops in `new_quiz_processing/100exam_hires_crops` to verify manually.")
    for item in ocr_items:
        rel = Path(item["crop"]).relative_to(BASE)
        lines.append("")
        lines.append(f"### {item['source_id']} ({item['quality']}, {item['line_count']} lines)")
        if item.get("duplicate_status") == "likely_duplicate":
            lines.append(f"Likely duplicate: `{item.get('best_match')}` score={item.get('best_score')}")
        if item.get("all_source_duplicate_status") == "likely_duplicate":
            lines.append(f"All-source likely duplicate: `{item.get('all_source_best_match')}` score={item.get('all_source_best_score')}")
        lines.append(f"Crop: `{rel.as_posix()}`")
        lines.append("```")
        lines.append(item["ocr_text"])
        lines.append("```")

    lines.append("")
    lines.append("## test_korka Title Index")
    for i, title in enumerate(test_titles, start=1):
        lines.append(f"- TK-{i:03d} p{title['page']}: {title['title']}")

    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main():
    ss6 = load_ss6()
    korsob = load_korsob()
    answer_cat = load_answer_category()
    unique, duplicates = dedupe(ss6 + korsob + answer_cat)
    ocr_items = load_ocr_items()
    ocr_dups = attach_duplicate_status(ocr_items, unique)
    ss5_items, ss5_dups = load_ss5_items(unique)
    exam02_items, exam02_dups = load_exam_items("exam02", EXAM02_OCR_DIR, "Exam02.pdf", "E02", unique)
    exam03_items, exam03_dups = load_exam_items("exam03", EXAM03_OCR_DIR, "Exam03.pdf", "E03", unique)
    test_titles = json.loads(TEST_KORKA_TITLES.read_text(encoding="utf-8")) if TEST_KORKA_TITLES.exists() else []
    source_groups = [
        ("SS5", ss5_items),
        ("Exam02", exam02_items),
        ("Exam03", exam03_items),
        ("100Exam", ocr_items),
    ]
    all_source_unique, all_source_dups, all_source_summary = compute_all_source_unique(unique, source_groups)

    build_docx(
        unique,
        duplicates,
        ss5_items,
        ss5_dups,
        exam02_items,
        exam02_dups,
        exam03_items,
        exam03_dups,
        ocr_items,
        ocr_dups,
        test_titles,
        all_source_unique,
        all_source_dups,
        all_source_summary,
    )
    build_markdown(
        unique,
        duplicates,
        ss5_items,
        ss5_dups,
        exam02_items,
        exam02_dups,
        exam03_items,
        exam03_dups,
        ocr_items,
        ocr_dups,
        test_titles,
        all_source_unique,
        all_source_dups,
        all_source_summary,
    )

    DUP_JSON.write_text(json.dumps(duplicates, ensure_ascii=False, indent=2), encoding="utf-8")
    SS5_DUP_JSON.write_text(json.dumps(ss5_dups, ensure_ascii=False, indent=2), encoding="utf-8")
    ref_dups = {
        "SS5": ss5_dups,
        "Exam02": exam02_dups,
        "Exam03": exam03_dups,
        "100Exam": ocr_dups,
    }
    REF_DUP_JSON.write_text(json.dumps(ref_dups, ensure_ascii=False, indent=2), encoding="utf-8")
    ALL_SOURCE_DUP_JSON.write_text(
        json.dumps(
            {
                "threshold": ALL_SOURCE_DUP_THRESHOLD,
                "all_source_unique_estimate": all_source_unique,
                "summary": all_source_summary,
                "duplicates": all_source_dups,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    source_inventory = {
        "verified_unique_count": len(unique),
        "duplicate_count": len(duplicates),
        "all_source_unique_estimate": all_source_unique,
        "all_source_duplicate_threshold": ALL_SOURCE_DUP_THRESHOLD,
        "all_source_likely_duplicate_count": len(all_source_dups),
        "all_source_summary": all_source_summary,
        "sources": {
            "SuperAI SS6_Foundation AI QUIZ.pdf": len(ss6),
            "SuperAI SS5_Foundation AI QUIZ.pdf OCR/crop appendix": len(ss5_items),
            "ข้อสอบชุด1.pdf": len(korsob),
            "[เฉลย] ข้อสอบ แยกหมวด.pdf": len(answer_cat),
            "Exam02.pdf OCR/crop appendix": len(exam02_items),
            "Exam03.pdf OCR/crop appendix": len(exam03_items),
            "100Exam_Lv1.pdf OCR/crop appendix": len(ocr_items),
            "test_korka_compressed.pdf title index": len(test_titles),
            "Quiz-92 score.docx/pdf": "answer-audit resource, not a separate full question set",
            "รวมข้อสอบSS6.txt": "text source for SS6 questions",
        },
        "ss5_likely_duplicate_count": len(ss5_dups),
        "exam02_likely_duplicate_count": len(exam02_dups),
        "exam03_likely_duplicate_count": len(exam03_dups),
        "100exam_likely_duplicate_count": len(ocr_dups),
        "known_source_corrections": ["K1-035", "K1-047"],
    }
    SOURCE_JSON.write_text(json.dumps(source_inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    print(DOCX_OUT)
    print(MD_OUT)
    print(SOURCE_JSON)
    print(ALL_SOURCE_DUP_JSON)
    print(f"verified_unique={len(unique)} duplicates={len(duplicates)} all_source_unique_estimate={all_source_unique} all_source_likely_dups={len(all_source_dups)}")
    print(f"ss5_items={len(ss5_items)} ss5_likely_dups={len(ss5_dups)} exam02_items={len(exam02_items)} exam02_likely_dups={len(exam02_dups)} exam03_items={len(exam03_items)} exam03_likely_dups={len(exam03_dups)} ocr_items={len(ocr_items)} ocr_likely_dups={len(ocr_dups)} test_titles={len(test_titles)}")


if __name__ == "__main__":
    main()
